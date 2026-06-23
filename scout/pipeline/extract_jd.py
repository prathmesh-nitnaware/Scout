"""
extract_jd.py — Extract and parse the actual job description from .docx
------------------------------------------------------------------------
Reads job_description.docx from the challenge directory.
Outputs artifacts/jd_profile.json with:
    {
        "role_title":      str,
        "role_confidence": float,
        "raw_text":        str,
        "required_skills": list[str],
        "preferred_skills": list[str],
        "responsibilities": list[str],
        "min_years":       int,
        "max_years":       int,
        "experience_text": str,
        "seniority":       str,
        "location":        str,
        "remote_ok":       bool,
        "word_count":      int,
        "char_count":      int,
    }

Also saves artifacts/job_description.txt for inspection.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

# MEDIUM ISSUE 6 FIX: Import centralized vocabulary
try:
    from scout.config.skills import SKILL_VOCAB, ROLE_PATTERNS
except ImportError:
    print("[JD] ERROR: scout.config.skills not found. Please ensure scout is in PYTHONPATH.")
    sys.exit(1)

CHALLENGE_DIR = (
    Path(__file__).resolve().parents[2]
    / "[PUB] India_runs_data_and_ai_challenge"
    / "India_runs_data_and_ai_challenge"
)
ARTIFACTS = Path(__file__).resolve().parents[2] / "artifacts"
DOCX_PATH = CHALLENGE_DIR / "job_description.docx"
TXT_OUT   = ARTIFACTS / "job_description.txt"
JSON_OUT  = ARTIFACTS / "jd_profile.json"

# Experience extraction patterns
EXP_PATTERNS = [
    re.compile(r"(\d+)\s*[-\u2013to]+\s*(\d+)\s*\+?\s*years?", re.IGNORECASE),
    re.compile(r"(?:at least|minimum|min\.?|>\s*)(\d+)\+?\s*years?", re.IGNORECASE),
    re.compile(r"(\d+)\+?\s*years?\s+(?:of\s+)?(?:experience|exp)", re.IGNORECASE),
]

SENIORITY_MAP = {
    "intern":   ["intern", "trainee", "fresher"],
    "junior":   ["junior", "jr", "entry-level", "entry level", "associate"],
    "mid":      ["mid-level", "mid level", "intermediate"],
    "senior":   ["senior", "sr.", "sr ", "experienced"],
    "lead":     ["lead", "tech lead", "team lead"],
    "staff":    ["staff", "principal", "distinguished"],
    "manager":  ["manager", "head of", "director"],
}

PREFERRED_MARKERS = re.compile(
    r"(preferred|nice to have|bonus|good to have|plus|desirable|optional)",
    re.IGNORECASE,
)
REMOTE_PATTERNS = re.compile(
    r"\b(remote|work from home|wfh|fully remote|distributed|hybrid)\b",
    re.IGNORECASE,
)


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract raw text from a .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(str(docx_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    return "\n".join(paragraphs)


def parse_role_title(text: str) -> tuple[str, float]:
    """
    Extract role title using word boundary regex patterns.
    Returns (Role Title, Confidence Score).
    """
    search_area = text[:1000].lower()

    # Priority 2: Sort longest roles first to prevent partial overlaps
    sorted_roles = sorted(ROLE_PATTERNS, key=len, reverse=True)

    for role in sorted_roles:
        # Priority 1: Use regex boundaries
        pattern = r"\b" + re.escape(role.lower()) + r"\b"
        if re.search(pattern, search_area):
            # Priority 5: Return 1.0 confidence on match
            return role.title(), 1.0

    return "Unknown", 0.0


def parse_experience(text: str) -> tuple[int, int, str]:
    """
    Extract min/max years of experience and the raw matched string.
    """
    for pat in EXP_PATTERNS:
        m = pat.search(text)
        if m:
            # Priority 3: Store raw match text
            raw_match = m.group(0).strip()
            groups = [g for g in m.groups() if g is not None]
            if len(groups) == 2:
                a, b = int(groups[0]), int(groups[1])
                return min(a, b), max(a, b), raw_match
            if len(groups) == 1:
                n = int(groups[0])
                return n, n + 3, raw_match
    return 0, 99, "Not specified"


def parse_seniority(text: str) -> str:
    lowered = text.lower()
    for level, keywords in SENIORITY_MAP.items():
        for kw in keywords:
            if kw in lowered:
                return level
    return "mid"


def extract_responsibilities(text: str) -> list[str]:
    """
    Priority 4: Extract bullet points specifically under a responsibilities section.
    """
    lines = text.splitlines()
    in_resp_section = False
    bullets = []
    
    for line in lines:
        lower_line = line.strip().lower()
        
        # State transitions
        if any(x in lower_line for x in ["responsibilities", "what you'll do", "what you will do", "duties"]):
            in_resp_section = True
            continue
        elif in_resp_section and any(x in lower_line for x in ["preferred", "nice to have", "requirements", "required skills", "qualifications"]):
            break  # We've hit the next major section header
            
        # Collect bullet points if in the correct state
        if in_resp_section:
            stripped = line.strip()
            if stripped.startswith(("-", "*", "•")):
                bullets.append(stripped.lstrip("-*• ").strip())
                
    return bullets


def split_jd_sections(text: str) -> tuple[str, str]:
    m = PREFERRED_MARKERS.search(text)
    if m:
        return text[:m.start()], text[m.start():]
    return text, ""


def extract_skills(text: str, vocab: list[str]) -> list[str]:
    found = []
    lowered = text.lower()
    for skill in vocab:
        pat = re.compile(r"\b" + re.escape(skill.lower()) + r"\b")
        if pat.search(lowered):
            found.append(skill)
    return found


def extract_location(text: str) -> str:
    pat = re.compile(
        r"(?:location|based in|office|city|work location)\s*[:\-\u2013]?\s*(.+)",
        re.IGNORECASE,
    )
    m = pat.search(text)
    if m:
        loc = m.group(1).strip().split("\n")[0].strip()
        loc = re.sub(r"\(.*?\)", "", loc).strip()
        return loc[:80]
    return "India"


def main() -> None:
    ARTIFACTS.mkdir(parents=True, exist_ok=True)

    if not DOCX_PATH.exists():
        print(f"[JD] ERROR: {DOCX_PATH} not found")
        raise SystemExit(1)

    print(f"[JD] Extracting text from: {DOCX_PATH.name}")
    raw_text = extract_text_from_docx(DOCX_PATH)
    print(f"[JD] Extracted {len(raw_text):,} chars, {len(raw_text.split()):,} words")

    TXT_OUT.write_text(raw_text, encoding="utf-8")
    print(f"[JD] Plain text saved: {TXT_OUT}")

    # Parse structure
    role_title, role_confidence = parse_role_title(raw_text)
    required_section, preferred_section = split_jd_sections(raw_text)
    required_skills  = extract_skills(required_section, SKILL_VOCAB)
    preferred_skills = [
        s for s in extract_skills(preferred_section, SKILL_VOCAB)
        if s not in required_skills
    ]
    responsibilities = extract_responsibilities(raw_text)
    min_years, max_years, experience_text = parse_experience(raw_text)
    seniority  = parse_seniority(raw_text)
    location   = extract_location(raw_text)
    remote_ok  = bool(REMOTE_PATTERNS.search(raw_text))

    profile = {
        "role_title":       role_title,
        "role_confidence":  role_confidence,
        "raw_text":         raw_text,
        "required_skills":  required_skills,
        "preferred_skills": preferred_skills,
        "responsibilities": responsibilities,
        "min_years":        min_years,
        "max_years":        max_years,
        "experience_text":  experience_text,
        "seniority":        seniority,
        "location":         location,
        "remote_ok":        remote_ok,
        "word_count":       len(raw_text.split()),
        "char_count":       len(raw_text),
    }

    JSON_OUT.write_text(json.dumps(profile, indent=2), encoding="utf-8")
    print(f"[JD] Profile saved: {JSON_OUT}")
    print(f"\n[JD] Parsed results:")
    print(f"  Role Title       : {role_title} (Confidence: {role_confidence})")
    print(f"  Required skills  : {len(required_skills)}  -> {required_skills[:8]}...")
    print(f"  Preferred skills : {len(preferred_skills)} -> {preferred_skills[:5]}")
    print(f"  Responsibilities : {len(responsibilities)} extracted")
    print(f"  Experience range : {min_years} - {max_years} years (Extracted from: '{experience_text}')")
    print(f"  Seniority        : {seniority}")
    print(f"  Location         : {location}")
    print(f"  Remote OK        : {remote_ok}")


if __name__ == "__main__":
    main()